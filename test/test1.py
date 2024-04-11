from docx import Document
import json


def read_docx(file_path):
    try:
        doc = Document(file_path)
        return doc
    except Exception as e:
        print("Error:", e)
        return None

if __name__ == "__main__":
    file_path = input("Enter the path of the Word document: ")

    # Read the Word document
    doc = read_docx(file_path)
    if doc:
        print("Document read successfully!")
    
    
    
#Function to parse the documents

def parse_questions(doc):
    questions = []
    current_question = None

    for i in range(len(doc.paragraphs) - 1):
        paragraph = doc.paragraphs[i]
        text = paragraph.text.strip()
       
        if text.startswith("Answer:"):
            if current_question:
                questions.append(current_question)
            current_question = None
            
        elif text:
            if not current_question:
                current_question = {"question": text}
            elif text.startswith("A."):
                
                
                if i + 3 < len(doc.paragraphs):
                    current_question["options"] = {
                        "A": text,
                        "B": doc.paragraphs[i + 1].text.strip(),
                        "C": doc.paragraphs[i + 2].text.strip(),
                        "D": doc.paragraphs[i + 3].text.strip(),
                    }
                else:
                    print("Error: Not enough paragraphs for options.")
            elif text.startswith("Answer:"):
                current_question["answer"] = text.split(":")[1].strip()
                
                
    # Append the last question if it exists.
    if current_question:
        questions.append(current_question)
        
    return questions

# Function to write parsed questions to a JSON file.

def write_to_json_file(data, file_path):
    try:
        with open(file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)
        print("Data written to", file_path)
    except Exception as e:
        print("Error:", e)
        
        
        
        

def print_json_file(file_path):
    
    try:
        with open(file_path, "r") as json_file:
            data = json.load(json_file)
            print(json.dumps(data, indent=4))
    except FileNotFoundError:
        print("Error: File not found.")
    except Exception as e:
        print("Error:", e)

doc = read_docx("questions.docx")
if doc:
    parsed_questions = parse_questions(doc)
    if parsed_questions:
        write_to_json_file(parsed_questions, "questions.json")
        print_json_file("questions.json")
    
    











