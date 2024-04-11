from docx import Document

def create_word_doc(file_path, questions):
    try:
        doc = Document()

        for i, q in enumerate(questions, start=1):
            doc.add_paragraph(f"{i}. {q['question']}")
            doc.add_paragraph(f"A. {q['options']['A']}")
            doc.add_paragraph(f"B. {q['options']['B']}")
            doc.add_paragraph(f"C. {q['options']['C']}")
            doc.add_paragraph(f"D. {q['options']['D']}")
            doc.add_paragraph(f"Answer: {q['answer']}")

        doc.save(file_path)
        print("Word document created successfully at", file_path)
        return file_path
    except Exception as e:
        print("Error:", e)
        return None

# Example questions
questions = [
    {
        "question": "What is the capital of France?",
        "options": {
            "A": "Berlin",
            "B": "Madrid",
            "C": "Paris",
            "D": "Lisbon"
        },
        "answer": "C"
    },
    {
        "question": "Who wrote Hamlet?",
        "options": {
            "A": "Mark Twain",
            "B": "William Shakespeare",
            "C": "Jane Austen",
            "D": "Charles Dickens"
        },
        "answer": "B"
    },
    
    {
        "question": "Who wrote Ramayana?",
        "options": {
            "A": "Valmiki",
            "B": "Jk Rowling",
            "C": "Jane Austen",
            "D": "Paulo Coelo"
        },
        "answer": "A"
    },
    
]

# Example usage:
create_word_doc("questions.docx", questions)
